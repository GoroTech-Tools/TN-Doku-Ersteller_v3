"""
core.py – Kernlogik für TN-Doku-Ersteller-Portable
Portierung der PowerShell-Funktionen nach Python ohne Office-COM-Abhängigkeit.
Verwendet: openpyxl (Excel), python-docx (Word), os/shutil (Dateisystem)
"""
import os
import shutil
import sys
from typing import Callable

import openpyxl
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Hilfsfunktionen
# ---------------------------------------------------------------------------

def get_app_dir() -> str:
    """Gibt das Basisverzeichnis der Anwendung zurück.
    Im EXE-Modus: Ordner der EXE-Datei.
    Im Entwicklungsmodus: Projekt-Wurzel (eine Ebene über src/).
    """
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def _log(msg: str, callback: Callable | None) -> None:
    if callback:
        callback(msg)
    else:
        print(msg)


# ---------------------------------------------------------------------------
# Schritt 1: Ordnerstruktur anlegen
# ---------------------------------------------------------------------------

def create_participant_folders(
    participants: list[dict],
    output_dir: str,
    log: Callable | None = None,
) -> str:
    """
    Legt 'Jahrgang XXXX' und darunter je einen Ordner pro Teilnehmer an.
    Gibt den Pfad zum Jahrgangsordner zurück.
    """
    if not participants:
        raise ValueError("Keine Teilnehmer in der CSV gefunden.")

    massnahme = participants[0].get('Maßnahme', '').strip()
    if len(massnahme) < 4:
        raise ValueError(
            f"'Maßnahme' im ersten Eintrag ist zu kurz oder leer: '{massnahme}'"
        )
    suffix = massnahme[-4:]
    jahrgang_name = f"Jahrgang {suffix}"
    jahrgang_path = os.path.join(output_dir, jahrgang_name)

    os.makedirs(jahrgang_path, exist_ok=True)
    _log(f"Jahrgangsordner: {jahrgang_path}", log)

    created = 0
    for p in participants:
        name = p.get('Name', '').strip()
        kuerzel = p.get('Maßnahmekürzel', '').strip()
        if not name or not kuerzel:
            _log(f"  Übersprungen (leere Felder): {p}", log)
            continue
        folder_name = f"{name} - {kuerzel}"
        folder_path = os.path.join(jahrgang_path, folder_name)
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
            created += 1
        _log(f"  Ordner: {folder_name}", log)

    _log(f"  → {created} Teilnehmer-Ordner erstellt.", log)
    return jahrgang_path


# ---------------------------------------------------------------------------
# Schritt 2: Ablagestruktur kopieren
# ---------------------------------------------------------------------------

def copy_template_structure(
    lists_dir: str,
    jahrgang_path: str,
    log: Callable | None = None,
) -> None:
    """
    Kopiert alle Unterordner aus lists_dir in jeden Teilnehmer-Ordner
    unterhalb von jahrgang_path.
    """
    if not os.path.isdir(lists_dir):
        raise FileNotFoundError(f"Ablagesystem-Ordner nicht gefunden: {lists_dir}")

    template_dirs = [
        d for d in os.scandir(lists_dir)
        if d.is_dir()
    ]
    if not template_dirs:
        _log("  Hinweis: Keine Unterordner im Ablagesystem gefunden.", log)
        return

    participant_folders = [
        d for d in os.scandir(jahrgang_path)
        if d.is_dir() and ' - ' in d.name
    ]
    if not participant_folders:
        raise FileNotFoundError(
            f"Keine Teilnehmer-Ordner in {jahrgang_path} gefunden."
        )

    for pf in participant_folders:
        for tmpl in template_dirs:
            dest = os.path.join(pf.path, tmpl.name)
            if not os.path.exists(dest):
                shutil.copytree(tmpl.path, dest)
        _log(f"  Struktur kopiert → {pf.name}", log)


# ---------------------------------------------------------------------------
# Schritt 3: Excel-Datei umbenennen (Muster → Teilnehmername)
# ---------------------------------------------------------------------------

EXCEL_TEMPLATE_NAME = "LEK- und ICDL-Ergebnisse Muster - KBM.xlsm"
LEK_SUBFOLDER = "LEK-Ergebnisse"


def rename_excel_files(
    jahrgang_path: str,
    log: Callable | None = None,
) -> None:
    """Benennt die XLSM-Musterdatei in jedem Teilnehmerordner um."""
    participant_folders = [
        d for d in os.scandir(jahrgang_path)
        if d.is_dir() and ' - ' in d.name
    ]
    renamed = 0
    for pf in participant_folders:
        lek_folder = os.path.join(pf.path, LEK_SUBFOLDER)
        if not os.path.isdir(lek_folder):
            _log(f"  Kein LEK-Ergebnisse-Ordner in: {pf.name}", log)
            continue
        template_path = os.path.join(lek_folder, EXCEL_TEMPLATE_NAME)
        if not os.path.isfile(template_path):
            continue
        new_name = f"LEK- und ICDL-Ergebnisse {pf.name}.xlsm"
        new_path = os.path.join(lek_folder, new_name)
        if os.path.exists(new_path):
            _log(f"  Bereits vorhanden: {new_name}", log)
            continue
        os.rename(template_path, new_path)
        _log(f"  Umbenannt → {new_name}", log)
        renamed += 1
    _log(f"  → {renamed} Excel-Dateien umbenannt.", log)


# ---------------------------------------------------------------------------
# Schritt 4: Excel-Tabellenblatt umbenennen (Muster → Teilnehmername)
# ---------------------------------------------------------------------------

def rename_excel_sheets(
    jahrgang_path: str,
    log: Callable | None = None,
) -> None:
    """Benennt das Tabellenblatt 'Muster' in der XLSM in den Teilnehmernamen um."""
    participant_folders = [
        d for d in os.scandir(jahrgang_path)
        if d.is_dir() and ' - ' in d.name
    ]
    processed = 0
    for pf in participant_folders:
        lek_folder = os.path.join(pf.path, LEK_SUBFOLDER)
        if not os.path.isdir(lek_folder):
            continue
        # Suche umbenannte XLSM
        xlsm_files = [
            f for f in os.scandir(lek_folder)
            if f.name.endswith('.xlsm') and 'Muster' not in f.name
        ]
        for xlsm in xlsm_files:
            try:
                wb = openpyxl.load_workbook(xlsm.path, keep_vba=True)
                if 'Muster' in wb.sheetnames:
                    # Nur der Name-Teil (vor " - ")
                    participant_name = pf.name.split(' - ')[0]
                    wb['Muster'].title = participant_name
                    wb.save(xlsm.path)
                    wb.close()
                    _log(f"  Blatt umbenannt → '{participant_name}' in {xlsm.name}", log)
                    processed += 1
                else:
                    wb.close()
            except Exception as e:
                _log(f"  Fehler beim Bearbeiten von {xlsm.name}: {e}", log)
    _log(f"  → {processed} Excel-Blätter umbenannt.", log)


# ---------------------------------------------------------------------------
# Schritt 5: Anwesenheitsliste erzeugen
# ---------------------------------------------------------------------------

def create_anwesenheitsliste(
    participants: list[dict],
    template_path: str,
    output_dir: str,
    log: Callable | None = None,
) -> None:
    """
    Erzeugt eine Anwesenheitsliste aus der Word-Vorlage.
    Füllt Kopfzeile und Tabelle mit Teilnehmerdaten.
    """
    if not os.path.isfile(template_path):
        raise FileNotFoundError(
            f"Anwesenheitsliste-Vorlage nicht gefunden: {template_path}"
        )

    massnahme = participants[0].get('Maßnahme', '').strip()
    suffix = massnahme[-4:] if len(massnahme) >= 4 else massnahme
    out_filename = f"Anwesenheitsliste KFL {suffix}.docx"
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, out_filename)

    doc = Document(template_path)

    # --- Kopfzeile setzen ---
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            if not para.text.strip():
                continue
            # Tab-Stop: dynamisch auf Inhaltsbreite setzen, damit "Datum:" am
            # rechten Rand endet und nicht darüber hinausschiebt.
            sec0 = doc.sections[0]
            content_w_twips = int(
                (sec0.page_width.pt - sec0.left_margin.pt - sec0.right_margin.pt) * 20
            )
            pPr = para._p.find(qn('w:pPr'))
            if pPr is not None:
                tabs_el = pPr.find(qn('w:tabs'))
                if tabs_el is not None:
                    for tab in tabs_el.findall(qn('w:tab')):
                        if tab.get(qn('w:val')) == 'right':
                            tab.set(qn('w:pos'), str(content_w_twips))
            # Runs leeren und neu befüllen
            for run in para.runs:
                run.text = ''
            runs = para.runs

            def _setr(idx: int, text: str, bold: bool | None = True) -> None:
                if idx < len(runs):
                    runs[idx].text = text
                    runs[idx].bold = bold
                else:
                    r = para.add_run(text)
                    r.bold = bold

            _setr(0, f"Anwesenheitsliste KFL {suffix}", True)
            _setr(1, '\t', True)
            _setr(2, 'Datum:', True)
            _setr(3, '________________', None)  # Schreibfeld
            break  # nur erster Absatz

    # --- Tabelle befüllen ---
    if doc.tables:
        table = doc.tables[0]

        # Leere Vorlage-Zeile(n) nach dem Header entfernen
        while len(table.rows) > 1:
            tr = table.rows[1]._tr
            table._tbl.remove(tr)

        # Teilnehmer einfügen
        for p in participants:
            name = p.get('Name', '').strip()
            kuerzel = p.get('Maßnahmekürzel', '').strip()
            row = table.add_row()
            row.cells[0].text = kuerzel
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
            row.cells[1].text = name

        # Zeilenhöhen berechnen: alle Datenzeilen füllen genau eine Seite
        sec0 = doc.sections[0]
        content_h_pt = (
            sec0.page_height.pt - sec0.top_margin.pt - sec0.bottom_margin.pt
        )
        overhead_pt = 35  # Schätzung für Tabellen-Kopfzeile + Abstände
        available_pt = content_h_pt - overhead_pt
        n = len(participants)
        row_h_pt = max(14.0, available_pt / n)
        row_h_twips = int(row_h_pt * 20)

        for row in table.rows[1:]:  # Nur Datenzeilen, nicht Header
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is None:
                trHeight = OxmlElement('w:trHeight')
                trPr.append(trHeight)
            trHeight.set(qn('w:val'), str(row_h_twips))
            trHeight.set(qn('w:hRule'), 'exact')
    else:
        _log("  Hinweis: Keine Tabelle in Vorlage gefunden – Liste ohne Tabelle.", log)

    doc.save(out_path)
    _log(f"  Anwesenheitsliste erstellt: {out_path}", log)


# ---------------------------------------------------------------------------
# Hauptablauf
# ---------------------------------------------------------------------------

def run_all(
    participants: list[dict],
    output_dir: str,
    lists_dir: str,
    log: Callable | None = None,
) -> str:
    """
    Führt alle Schritte aus und gibt den Pfad zum Jahrgangsordner zurück.
    """
    _log("Schritt 1/5: Ordnerstruktur anlegen …", log)
    jahrgang_path = create_participant_folders(participants, output_dir, log)

    _log("Schritt 2/5: Ablagestruktur kopieren …", log)
    copy_template_structure(lists_dir, jahrgang_path, log)

    _log("Schritt 3/5: Excel-Dateien umbenennen …", log)
    rename_excel_files(jahrgang_path, log)

    _log("Schritt 4/5: Excel-Tabellenblätter umbenennen …", log)
    rename_excel_sheets(jahrgang_path, log)

    _log("Schritt 5/5: Anwesenheitsliste erstellen …", log)
    anwesenheit_template = os.path.join(lists_dir, 'Anwesenheitsliste.docx')
    create_anwesenheitsliste(
        participants,
        anwesenheit_template,
        output_dir,
        log=log,
    )

    _log(f"\n✓ Fertig! Ergebnis: {jahrgang_path}", log)
    return jahrgang_path
